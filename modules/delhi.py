import pdfplumber
import re
import os
import docx2pdf
import logging
from docx import Document

# Configure the logger to append to the log file
logging.basicConfig(filename='delhi_process_log.log', filemode='a', 
                    format='%(asctime)s - %(levelname)s - %(message)s', 
                    level=logging.INFO)

def extract_delhi_dsi(pdf_path):
    def convert_to_pdf(input_file):
        """Converts .doc or .docx files to .pdf for processing."""
        try:
            logging.info(f"Starting conversion of {input_file} to PDF.")
            if not os.path.isfile(input_file):
                logging.error(f"File {input_file} does not exist.")
                return None

            file_extension = os.path.splitext(input_file)[1].lower()
            base_name = os.path.splitext(input_file)[0]

            if file_extension == '.doc':
                logging.info(f"Converting .doc to .docx for file: {input_file}")
                document = Document()
                document.LoadFromFile(input_file)
                docx_file = f"{base_name}.docx"
                document.SaveToFile(docx_file, FileFormat.Docx2016)
                document.Close()
                input_file = docx_file
                logging.info(f".doc file converted to .docx: {docx_file}")

            if input_file.endswith('.docx'):
                pdf_file = f"{base_name}.pdf"
                docx2pdf.convert(input_file, pdf_file)
                logging.info(f".docx file converted to .pdf: {pdf_file}")
                return pdf_file  # Return the path to the new PDF file

            if input_file.endswith('.pdf'):
                logging.info(f"File is already a PDF: {input_file}")
                return input_file  # If already a PDF, return as is
        except Exception as e:
            logging.error(f"Error converting file {input_file}: {e}")
            return None

    def is_page_number(line):
        result = line.strip().isdigit() or line.strip().lower().startswith(('page', 'p age'))
        logging.debug(f"Checking if line is page number '{line}': {result}")
        return result

    def extract_filtered_text_to_variable(pdf_path):
        all_text = []
        try:
            logging.info(f"Extracting and filtering text from PDF: {pdf_path}")
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        logging.debug(f"Extracting text from page {page_num}")
                        filtered_lines = [line for line in text.split('\n') if not is_page_number(line)]
                        filtered_text = '\n'.join(filtered_lines)
                        all_text.append(filtered_text)
            logging.info(f"Text extracted and filtered successfully from {pdf_path}")
            return '\n\n'.join(all_text)
        except Exception as e:
            logging.error(f"Error extracting text from PDF {pdf_path}: {e}")
            return ""

    def clean_text(text):
        logging.info("Cleaning text by removing sensitive words.")
        cleaned_text = re.sub(r'\bSECRET\b', '', text, flags=re.IGNORECASE)
        cleaned_text = re.sub(r'\bCOMMUNAL\b', '', cleaned_text, flags=re.IGNORECASE)
        parts = re.split(r'\n\*+\n', cleaned_text, maxsplit=1)
        logging.debug("Text cleaned successfully.")
        return parts[1] if len(parts) > 1 else ""

    def clean_headings(text):
        logging.info("Cleaning headings in the text.")
        pattern = r'\n((?:[A-Z]{4,}\.?)|(?:YOUTH (?:AND|&) STUDENTS))\n'
        matches = re.findall(pattern, text)
        for match in matches:
            cleaned_match = re.sub(r'[^A-Za-z0-9 ]', '', match)
            text = re.sub(rf'(?<=\n){re.escape(match)}(?=\n)', cleaned_match, text)
        logging.debug("Headings cleaned successfully.")
        return text

    def find_headings(text):
        logging.info("Finding headings in the text.")
        pattern = r'\n((?:[A-Z]{4,}\.?)|(?:YOUTH (?:AND|&) STUDENTS))\n'
        matches = re.findall(pattern, text)
        logging.debug(f"Headings found: {matches}")
        return [match.rstrip('.') for match in matches]

    def is_heading(line):
        result = line.isupper() and not re.search(r'[.!?;]$', line)
        logging.debug(f"Checking if line is heading '{line}': {result}")
        return result

    def add_heading_to_entries_with_continuous_numbering(text, restart_numbering=False):
        logging.info("Adding headings to entries with continuous numbering.")
        lines = text.split('\n')
        updated_lines, current_heading, continuous_number = [], None, 1
        for line in lines:
            stripped_line = line.strip()
            if is_heading(stripped_line):
                current_heading = stripped_line
                if restart_numbering:
                    continuous_number = 1
                continue
            match = re.match(r'^(\d+)\.\s*(.*)', stripped_line)
            if match:
                entry_number, entry_text = match.groups()
                if int(entry_number) == continuous_number:
                    updated_lines.append(f"{current_heading}-{continuous_number}.{entry_text}")
                    continuous_number += 1
                else:
                    updated_lines.append(line)
            else:
                updated_lines.append(line)
        logging.debug("Headings added successfully to continuous entries.")
        return '\n'.join(updated_lines)

    def split_text_by_keywords(text, headings):
        logging.info("Splitting text by keywords.")
        text = re.sub(r'\n', ' ', text)
        pattern = r'(' + '|'.join([re.escape(word) for word in headings]) + r')(.*?)(?=' + '|'.join([re.escape(word) for word in headings]) + r'|\Z)'
        matches = re.findall(pattern, text, re.DOTALL)
        logging.debug(f"Text split by keywords, total sections: {len(matches)}")
        return [[f"{heading} {content.strip()}"] for heading, content in matches]

    # Convert to PDF if necessary
    logging.info(f"Starting file check and conversion for: {pdf_path}")
    checked = convert_to_pdf(pdf_path)
    if not checked:
        logging.error("The file could not be converted to PDF or does not exist.")
        raise ValueError("The file could not be converted to PDF or does not exist.")

    extracted_text = extract_filtered_text_to_variable(checked)
    cleaned_text = clean_text(extracted_text)
    cleaned_headings = clean_headings(cleaned_text)
    heading_added = add_heading_to_entries_with_continuous_numbering(cleaned_headings)
    headings = find_headings(cleaned_headings)
    final_output = split_text_by_keywords(heading_added, headings)

    logging.info("Text extraction and processing completed successfully.")
    return final_output
